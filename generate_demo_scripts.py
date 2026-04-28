"""
Generate BoostCTC Voice Agent Demo Scripts as a DOCX file.
Run: python generate_demo_scripts.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colour palette ────────────────────────────────────────────────────────────
BRAND_PURPLE  = RGBColor(0x6C, 0x3A, 0xF5)   # BoostCTC primary
MIRA_BLUE     = RGBColor(0x15, 0x6B, 0xC9)   # Mira speech bubbles
USER_GRAY     = RGBColor(0x2C, 0x2C, 0x2C)   # User speech
STAGE_DIR_CLR = RGBColor(0x70, 0x70, 0x70)   # Stage directions
ACCENT_TEAL   = RGBColor(0x0A, 0x8A, 0x8A)   # Section labels
WHITE         = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG      = RGBColor(0xF4, 0xF0, 0xFF)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, fill_hex: str):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

# ── Helper: add a horizontal rule ────────────────────────────────────────────
def add_hr(document):
    p = document.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0B8F0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)
    return p

# ── Helper: cover title block ─────────────────────────────────────────────────
def add_cover(document):
    # big title
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("BoostCTC Voice Agent")
    set_font(r, size=28, bold=True, color=BRAND_PURPLE)

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(4)
    r2 = p2.add_run("Demo Conversation Scripts")
    set_font(r2, size=18, bold=False, color=MIRA_BLUE)

    p3 = document.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run("Three persona walkthroughs — illustrating Socratic coaching, "
                     "persona-aware exercises, and personalised advocacy")
    set_font(r3, size=10, italic=True, color=STAGE_DIR_CLR)

    add_hr(document)

    # Legend
    p4 = document.add_paragraph()
    p4.paragraph_format.space_before = Pt(10)
    p4.paragraph_format.space_after  = Pt(2)
    r4 = p4.add_run("How to read these scripts")
    set_font(r4, size=11, bold=True, color=BRAND_PURPLE)

    legend_items = [
        ("MIRA",         MIRA_BLUE,     "AI voice agent — spoken aloud by Vapi"),
        ("USER",         USER_GRAY,     "Visitor / returning user — spoken response"),
        ("[Stage note]", STAGE_DIR_CLR, "Italicised direction — not spoken"),
    ]
    for label, clr, desc in legend_items:
        pl = document.add_paragraph()
        pl.paragraph_format.left_indent = Cm(0.5)
        pl.paragraph_format.space_after = Pt(2)
        rl = pl.add_run(f"{label}  ")
        set_font(rl, size=10, bold=True, color=clr)
        rd = pl.add_run(desc)
        set_font(rd, size=10, italic=True, color=STAGE_DIR_CLR)

    add_hr(document)

# ── Helper: section heading ───────────────────────────────────────────────────
def add_script_header(document, number: str, title: str, subtitle: str, tags: list[str]):
    document.add_page_break()

    # Number pill
    p0 = document.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p0.paragraph_format.space_before = Pt(6)
    p0.paragraph_format.space_after  = Pt(2)
    r0 = p0.add_run(f"  SCRIPT {number}  ")
    set_font(r0, size=9, bold=True, color=WHITE)
    r0.font.highlight_color = None
    shade_paragraph(p0, "6C3AF5")

    # Title
    pt = document.add_paragraph()
    pt.paragraph_format.space_after = Pt(2)
    rt = pt.add_run(title)
    set_font(rt, size=16, bold=True, color=BRAND_PURPLE)

    # Subtitle
    ps = document.add_paragraph()
    ps.paragraph_format.space_after = Pt(8)
    rs = ps.add_run(subtitle)
    set_font(rs, size=10, italic=True, color=STAGE_DIR_CLR)

    # Tags
    pt2 = document.add_paragraph()
    pt2.paragraph_format.space_after = Pt(10)
    for tag in tags:
        rt2 = pt2.add_run(f"  {tag}  ")
        set_font(rt2, size=9, bold=True, color=ACCENT_TEAL)
        rt2.font.highlight_color = None

    add_hr(document)

# ── Helper: add a dialogue line ───────────────────────────────────────────────
def add_line(document, speaker: str, text: str, is_stage: bool = False):
    p = document.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.space_before = Pt(0)

    if is_stage:
        r = p.add_run(f"[{text}]")
        set_font(r, size=10, italic=True, color=STAGE_DIR_CLR)
        return

    color = MIRA_BLUE if speaker == "MIRA" else USER_GRAY
    rl = p.add_run(f"{speaker}:  ")
    set_font(rl, size=10, bold=True, color=color)
    rd = p.add_run(text)
    set_font(rd, size=10.5, color=color)

def add_phase_label(document, label: str):
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f"▸ {label}")
    set_font(r, size=9, bold=True, color=ACCENT_TEAL)
    r.font.italic = True

# ═════════════════════════════════════════════════════════════════════════════
#  COVER
# ═════════════════════════════════════════════════════════════════════════════
add_cover(doc)


# ═════════════════════════════════════════════════════════════════════════════
#  SCRIPT 1 — Early Career / Interview Prep / System Design
# ═════════════════════════════════════════════════════════════════════════════
add_script_header(
    doc,
    number="1",
    title="Early-Career Engineer — Interview Prep",
    subtitle="Persona: Priya, 2 years into software engineering, Google interview in 3 weeks, struggles with system design rounds",
    tags=["NEW VISITOR", "CAREER STAGE: EARLY", "FOCUS: INTERVIEW PREP / SYSTEM DESIGN"],
)

add_phase_label(doc, "PHASE 1 — Elevator Pitch  (delivered by Vapi as firstMessage)")

add_line(doc, "MIRA",
    "Hey, welcome to BoostCTC! Here's the thing most people don't realise — it's not about being smart, "
    "it's about being able to think clearly and communicate that thinking under pressure. That's the gap that "
    "holds careers back and makes high-stakes moments feel harder than they need to be. We close that gap with "
    "daily AI-coached exercises — Socratic scenarios, critical thinking drills, communication practice — and in "
    "this very call I'm going to walk you through one so you can feel it firsthand. "
    "But first, quick question: are you here for yourself, for your kid, or for a classroom?")

add_line(doc, "PRIYA", "For myself.")

add_phase_label(doc, "PHASE 2 — Discovery  (engagement_greeting)")

add_line(doc, "MIRA",
    "Got it — what's the main thing you're hoping to work on? Like is it something specific: "
    "nailing interviews, getting your ideas across at work, or something else?")

add_line(doc, "PRIYA",
    "I have a big tech interview coming up in a few weeks and I'm really struggling with system design rounds. "
    "I can code fine, but when they ask me to design a whole system from scratch I kind of freeze. "
    "I don't know where to start and I can't seem to organise my thinking out loud.")

add_line(doc, "MIRA",
    "That makes total sense — system design is one of those rounds where technical knowledge alone doesn't cut it. "
    "You need to think out loud, structure your reasoning in real time, and communicate tradeoffs clearly while "
    "someone is watching. That's actually very much what we help people build. "
    "And you're currently working as a software engineer?")

add_line(doc, "PRIYA", "Yeah, about two years in at a startup.")

add_line(doc, "MIRA",
    "Perfect. We have a great exercise for exactly this — I'll actually walk you through one in a minute. "
    "But first — when you say you freeze, is it more that you don't know the right answer, or that you "
    "have ideas but can't get them out fast enough?")

add_line(doc, "PRIYA",
    "Honestly the second one. I usually have a rough idea but then I second-guess myself and start rambling "
    "and I can tell the interviewer is losing confidence in me.")

add_line(doc, "MIRA",
    "Got it — so the thinking is there, the communication under pressure is what breaks down. "
    "That's a really specific thing to train, and it's very fixable. Alright, let's do this.")

add_phase_label(doc, "PHASE 3 — Value Exploration  (value_exploration)")

add_line(doc, "MIRA",
    "So what we train at BoostCTC is the rhythm of daily practice, AI feedback, and continuous progress — "
    "think of it like a gym for your thinking and communication muscle. "
    "The Socratic exercises in particular are built around thinking out loud: you're given a scenario, "
    "you reason through it, and the AI gives you feedback not just on whether you're right, "
    "but on how clearly you structured your thinking. For interview prep, especially system design, "
    "that 'reasoning out loud' muscle is exactly what separates people who pass from people who know the material "
    "but don't get the offer. Does that resonate with what you're running into?")

add_line(doc, "PRIYA",
    "Yeah, exactly. My interviewer last time said I had good ideas but my answer was hard to follow.")

add_line(doc, "MIRA",
    "That's the gap — and it's so coachable. Alright, want to feel what one of these exercises actually "
    "feels like? I'll give you a live sample right now.")

add_phase_label(doc, "PHASE 4 — Socratic Exercise  (socratic_taste)")

add_line(doc, "MIRA",
    "Here's a scenario. You're midway through a system design interview — you've sketched out a reasonable "
    "architecture for a URL shortener. Your interviewer says: 'How would this system handle ten times the "
    "current load?' You have about 60 seconds. What's the first thing you say?")

add_line(doc, "PRIYA",
    "Um... I'd probably say we should add more servers? And maybe put a load balancer in front?")

add_line(doc, "MIRA",
    "That's exactly the right direction — horizontal scaling is the correct first instinct. "
    "Here's what our exercises train on top of that: the way you frame it. "
    "Compare these two answers. 'Add more servers' versus "
    "'I'd scale horizontally by adding application servers behind a load balancer — "
    "that buys us capacity, but it also means we need to make the service fully stateless, "
    "which has implications for session management.' "
    "Same knowledge. Completely different signal to the interviewer. "
    "One sounds like a guess, one sounds like an engineer who's thought through the consequences. "
    "That's literally what BoostCTC trains every single day — not the knowledge, the articulation under pressure. "
    "Honestly, based on what you just said, I think you'd get a lot out of this. Should I get you set up?")

add_line(doc, "PRIYA", "Yeah, absolutely. That's exactly the gap I need to close.")

add_phase_label(doc, "PHASE 5 — Lead Capture  (lead_capture)")

add_line(doc, "MIRA",
    "Awesome — there's a quick form on your screen right now. "
    "Just type your name and email there and we'll get you all set!")

add_line(doc, "", "", is_stage=True)
doc.paragraphs[-1].runs[0].text = "Lead capture form appears on screen. Priya types her name and email."

add_line(doc, "MIRA",
    "You're all set, Priya! Check your inbox for the welcome email. "
    "Given what you're working on, I'd start with the critical thinking passage set — "
    "it's built to help you structure complex reasoning quickly under time pressure. "
    "Good luck with the interview. Go get it!")


# ═════════════════════════════════════════════════════════════════════════════
#  SCRIPT 2 — Mid Career / Critical Thinking + Communication
# ═════════════════════════════════════════════════════════════════════════════
add_script_header(
    doc,
    number="2",
    title="Mid-Career Product Manager — Critical Thinking & Communication",
    subtitle="Persona: Marcus, 8 years in B2B SaaS product management, wants to build stronger data-driven arguments and persuade cross-functional stakeholders",
    tags=["NEW VISITOR", "CAREER STAGE: MID", "FOCUS: CRITICAL THINKING + COMMUNICATION"],
)

add_phase_label(doc, "PHASE 1 — Elevator Pitch  (firstMessage)")

add_line(doc, "MIRA",
    "Hey, welcome to BoostCTC! Here's the thing most people don't realise — it's not about being smart, "
    "it's about being able to think clearly and communicate that thinking under pressure. "
    "That's the gap that holds careers back, causes students to underperform, and makes meetings "
    "feel unproductive. We close that gap with daily AI-coached exercises — Socratic scenarios, "
    "critical thinking drills, communication practice — and in this very call, I'm going to walk you through "
    "one so you can feel it firsthand, not just hear about it. "
    "But first, quick question: are you here for yourself, for your kid, or for a classroom?")

add_line(doc, "MARCUS", "For myself.")

add_phase_label(doc, "PHASE 2 — Discovery  (engagement_greeting)")

add_line(doc, "MIRA",
    "Got it — what's the main thing you're hoping to work on? "
    "Is it something specific or more general?")

add_line(doc, "MARCUS",
    "I'm a product manager and I keep getting feedback that I need to be more data-driven in how I make "
    "arguments. Like I'll have a strong gut feeling about the right direction, but I struggle to build "
    "a case that actually convinces engineers and stakeholders who think very differently from me.")

add_line(doc, "MIRA",
    "That's a really honest insight — and it's one of the most common friction points for experienced PMs. "
    "The challenge is bridging the gap between strategic intuition and making it land for people "
    "who need logic trails, not just conviction. Are you leading a team currently?")

add_line(doc, "MARCUS",
    "Yeah, I've been a PM for about eight years, managing a team for the last three. Series B company.")

add_line(doc, "MIRA",
    "Got it. So you've been in enough rooms to know your read is usually right — "
    "the gap is in the translation. We have a really good exercise for exactly this. "
    "Let me walk you through one.")

add_phase_label(doc, "PHASE 3 — Value Exploration  (value_exploration)")

add_line(doc, "MIRA",
    "At BoostCTC, the core training loop is daily practice, AI feedback, and continuous progress. "
    "For someone at your stage, the exercises that matter most are the critical thinking passages — "
    "where you reason through a scenario and explain your logic out loud — and the communication exercises, "
    "where you take a complex idea and the AI gives you specific feedback on how clearly it landed. "
    "The thing that makes it different is the AI doesn't just say 'good job' — it tells you exactly "
    "where your argument breaks down and why. For a PM trying to move a room, that kind of specific "
    "feedback is hard to get anywhere else. Does the reasoning-out-loud part feel like the bigger gap, "
    "or is it more the communication and framing side?")

add_line(doc, "MARCUS",
    "Honestly both, but if I had to pick one — I think I can do the analysis, "
    "it's more that I can't always get people to follow my reasoning in the moment.")

add_line(doc, "MIRA",
    "Communication under pressure. Got it. That's exactly what the Socratic exercises train. "
    "Want to try one right now?")

add_line(doc, "MARCUS", "Sure.")

add_phase_label(doc, "PHASE 4 — Socratic Exercise  (socratic_taste)")

add_line(doc, "MIRA",
    "Here's the scenario. You're in a strategy review and your team presents an AI-generated market analysis "
    "recommending that you cut your enterprise tier entirely and go all-in on SMB. "
    "The numbers look clean. Three people in the room are already nodding. "
    "Your gut says something is off — but you haven't fully worked out what yet. "
    "You have about two minutes to respond. What do you say?")

add_line(doc, "MARCUS",
    "I'd want to understand the assumptions behind the data first. "
    "Like, which customer segments were actually included in this? "
    "And I'd push on the long-term revenue implications — "
    "not just the short-term numbers that make the SMB case look clean.")

add_line(doc, "MIRA",
    "That's really good — you went straight to the assumptions before reacting, "
    "and you immediately introduced a time horizon the model probably hadn't accounted for. "
    "That's exactly the kind of critical thinking that separates strong PMs from great ones. "
    "What BoostCTC trains on top of that is the habit of doing it automatically and out loud — "
    "because in that room you don't have time to think silently first. "
    "You need to be able to say 'before we go further I want to stress-test two assumptions' "
    "in a way that leads the room rather than just raising doubt. "
    "The difference between those two things is why some people are seen as strategic and others aren't, "
    "even when they're thinking the same thoughts. I think you'd genuinely get a lot out of this. "
    "Want to get set up?")

add_line(doc, "MARCUS",
    "Definitely. This is exactly the muscle I've been trying to build for years.")

add_phase_label(doc, "PHASE 5 — Lead Capture  (lead_capture)")

add_line(doc, "MIRA",
    "Love it. There's a quick form on your screen — just type your name and email "
    "and we'll get you started.")

add_line(doc, "", "", is_stage=True)
doc.paragraphs[-1].runs[0].text = "Lead capture form appears on screen. Marcus types his name and email."

add_line(doc, "MIRA",
    "You're all set, Marcus! I'd jump straight into the critical thinking passage set — "
    "specifically the 'challenge the assumption' exercises. Given what you just demonstrated in that scenario, "
    "you'll progress fast. The AI feedback will sharpen exactly the gap you described. "
    "Welcome to BoostCTC!")


# ═════════════════════════════════════════════════════════════════════════════
#  SCRIPT 3 — Returning User / Existing Advocate
# ═════════════════════════════════════════════════════════════════════════════
add_script_header(
    doc,
    number="3",
    title="Returning User — Personalised Advocacy",
    subtitle="Persona: Alex, consultant, 5 weeks on platform. Strong consistency (12-day streak, 34 passages) but low critical thinking score (2.8/5). Communication score 3.6/5.",
    tags=["EXISTING USER", "GAP: CRITICAL THINKING", "SCENARIO: CONSULTING CONTEXT"],
)

add_phase_label(doc, "PHASE 1 — Welcome Back  (advocacy_greeting)")

add_line(doc, "MIRA",
    "Hey, welcome back! I'm Mira — I can actually see how your recent sessions have been going. "
    "Got a minute to chat about it?")

add_line(doc, "ALEX", "Hey Mira, yeah sure, what's up?")

add_line(doc, "MIRA",
    "Great — and honestly, first thing I want to say is the streak is impressive. "
    "Twelve days straight is real commitment, and I can see you've done 34 passages. "
    "That consistency puts you in the top tier of users at your stage. "
    "Let's take a look at how things are going.")

add_phase_label(doc, "PHASE 2 — Performance Review  (performance_review)")

add_line(doc, "MIRA",
    "So here's the picture. Your communication score is sitting at 3.6 out of 5 — "
    "that's genuinely solid, you've clearly been putting work in there. "
    "Your critical thinking score is at 2.8, which is lower, and your MCQ accuracy is around 58 percent. "
    "Looking at those numbers, what do you think is your biggest opportunity to grow?")

add_line(doc, "ALEX",
    "Yeah, critical thinking for sure. I feel it when I'm reading complex passages — "
    "I can understand them but when I try to answer the analysis questions I kind of go blank. "
    "Or I'll write something and then delete it and start over three times.")

add_line(doc, "MIRA",
    "That's a really specific observation — you understand the content but the analytical step breaks down. "
    "Let me dig into that a bit.")

add_phase_label(doc, "PHASE 3 — Gap Deepdive  (gap_deepdive)")

add_line(doc, "MIRA",
    "When you say you go blank — is it that you don't know where to start, "
    "or that you start and then lose confidence halfway through?")

add_line(doc, "ALEX",
    "Definitely the second one. I'll have an initial take and then I'll start doubting it "
    "and end up changing my answer — and I usually change it to the wrong one. "
    "Like my gut is right but I talk myself out of it.")

add_line(doc, "MIRA",
    "That's useful. So the issue isn't the thinking — it's the confidence in the thinking. "
    "Your instinct is usually right but you don't trust it enough to hold it. "
    "Does that show up anywhere outside the platform? Like in your work?")

add_line(doc, "ALEX",
    "Honestly, yeah. I'm in consulting and in client meetings I'll have a strong point of view, "
    "but the moment a senior partner or the client pushes back even slightly "
    "I immediately back down — even when I'm pretty sure I'm right. "
    "It's cost me credibility a few times.")

add_line(doc, "MIRA",
    "That's a really honest read — and it takes self-awareness to see that pattern clearly. "
    "So the through-line is: strong initial judgment, but second-guessing under any kind of pressure — "
    "whether it's a passage or a partner in the room. That's exactly what I want to address.")

add_phase_label(doc, "PHASE 4 — Personalised Plan  (personalized_nudge)")

add_line(doc, "MIRA",
    "Here's what I'd suggest specifically for you. "
    "There's a format in the critical thinking exercises called 'defend your reasoning' — "
    "you commit to an answer, and then the AI pushes back on it, asks you to justify it, "
    "even challenges you with counter-arguments. "
    "The whole design is to train the muscle of holding your position long enough "
    "to actually evaluate whether the pushback has merit, rather than just caving reflexively. "
    "Because right now you're abandoning positions that are often right — "
    "and the gap between your gut judgment and your confidence in it is exactly what that exercise closes. "
    "Even three of those sessions a week, ten minutes each, and you'll start to notice the difference "
    "within the first two weeks. Does that feel like the right direction?")

add_line(doc, "ALEX",
    "Yeah, absolutely. I had no idea there was something that specific. "
    "That's basically describing the exact problem I have.")

add_line(doc, "MIRA",
    "The specificity is the point — generic exercises won't move this needle for you. "
    "So what do you say — want to commit to three of those sessions this week?")

add_line(doc, "ALEX", "Yeah, I'll do it. Three sessions this week for sure.")

add_phase_label(doc, "PHASE 5 — Wrap Up  (engagement_wrapup)")

add_line(doc, "MIRA",
    "You've got this, Alex. That 'defend your reasoning' practice is going to pay off faster than you think. "
    "And next time a partner pushes back in a client meeting, you'll have the muscle to pause, "
    "hold your ground, and evaluate before you respond. "
    "That's a completely different kind of credibility. "
    "Go make this week count — see you next time!")

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = "/Users/sughosh/Downloads/BoostCTC Voice Agent/BoostCTC_Demo_Scripts.docx"
doc.save(out_path)
print(f"✓ Saved: {out_path}")
